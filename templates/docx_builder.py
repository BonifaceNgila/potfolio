from io import BytesIO

from utils.converters import normalize_education_record, normalize_project_record

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    DOCX_AVAILABLE = True
except ModuleNotFoundError:
    docx = None
    Pt = None
    RGBColor = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    DOCX_AVAILABLE = False


if DOCX_AVAILABLE:
    ACCENT = RGBColor(30, 58, 95)
    TEXT = RGBColor(31, 41, 55)
    MUTED = RGBColor(85, 98, 112)
else:
    ACCENT = None
    TEXT = None
    MUTED = None


def _clean_docx_text(value: object) -> str:
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 4, line_spacing: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _add_paragraph(doc, text: object = "", *, bold: bool = False, italic: bool = False, size: int = 10, color=None, style=None):
    paragraph = doc.add_paragraph(style=style)
    _set_paragraph_spacing(paragraph)
    run = paragraph.add_run(_clean_docx_text(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or TEXT
    return paragraph


def _add_bullet(doc, text: object, *, size: int = 10, indent: float = 0.2):
    paragraph = _add_paragraph(doc, text, size=size, style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def build_docx(cv: dict, template: str) -> bytes:
    if not DOCX_AVAILABLE:
        return b""

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = TEXT

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(name_para, after=2, line_spacing=1.0)
    name_run = name_para.add_run(_clean_docx_text(cv.get("full_name", "")))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ACCENT

    headline_para = doc.add_paragraph()
    headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(headline_para, after=8, line_spacing=1.0)
    headline_run = headline_para.add_run(_clean_docx_text(cv.get("headline", "")))
    headline_run.font.size = Pt(12)
    headline_run.font.color.rgb = MUTED

    contact_parts = []
    if cv.get("location"):
        contact_parts.append(("Location", cv.get("location")))
    if cv.get("phone"):
        contact_parts.append(("Phone", cv.get("phone")))
    if cv.get("email"):
        contact_parts.append(("Email", cv.get("email")))
    if cv.get("linkedin"):
        contact_parts.append(("LinkedIn", cv.get("linkedin")))
    if cv.get("github"):
        contact_parts.append(("GitHub", cv.get("github")))

    if contact_parts:
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for label, value in contact_parts:
            cells = table.add_row().cells
            cells[0].width = Inches(1.0)
            cells[1].width = Inches(5.8)
            for cell in cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            label_para = cells[0].paragraphs[0]
            _set_paragraph_spacing(label_para, after=1)
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = ACCENT
            value_para = cells[1].paragraphs[0]
            _set_paragraph_spacing(value_para, after=1)
            value_run = value_para.add_run(_clean_docx_text(value))
            value_run.font.size = Pt(8.5)
            value_run.font.color.rgb = TEXT

    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, after=4)

    def add_heading(text):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=8, after=4, line_spacing=1.0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(_clean_docx_text(text).upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = ACCENT
        return p

    if cv.get("profile_summary"):
        add_heading("Profile")
        p = _add_paragraph(doc, cv.get("profile_summary"), size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if cv.get("core_competencies"):
        add_heading("Core Competencies")
        for item in cv.get("core_competencies"):
            if str(item).strip():
                _add_bullet(doc, item, size=9.5)

    if cv.get("experience"):
        add_heading("Professional Experience")
        for exp in cv.get("experience"):
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=2)
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run(_clean_docx_text(exp.get("role", "")))
            r1.bold = True
            r1.font.size = Pt(10.5)
            r1.font.color.rgb = TEXT
            meta = p.add_run(
                f" - {_clean_docx_text(exp.get('organization', ''))} | {_clean_docx_text(exp.get('period', ''))}"
            )
            meta.font.size = Pt(9.5)
            meta.font.color.rgb = MUTED

            for bullet in exp.get("bullets", []):
                if str(bullet).strip():
                    bp = _add_bullet(doc, bullet, size=9.5)
                    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if cv.get("education"):
        add_heading("Education")
        for item in cv.get("education"):
            record = normalize_education_record(item)
            course = record.get("course", "")
            institution = record.get("institution", "")
            timeline = record.get("timeline", "")

            parts = [p for p in [course, institution] if p]
            if timeline:
                parts.append(f"({timeline})")

            if parts:
                _add_bullet(doc, " - ".join(parts), size=9.5)

    if cv.get("projects"):
        add_heading("Projects")
        for item in cv.get("projects"):
            record = normalize_project_record(item)
            name = record.get("name", "")
            description = record.get("description", "")
            technologies = record.get("technologies", "")
            link = record.get("link", "")
            if not name.strip() and not description.strip():
                continue
            if name:
                p = _add_bullet(doc, name, size=9.5)
                p.paragraph_format.keep_with_next = bool(description or technologies or link)
                p.runs[0].bold = True
            if description:
                p = _add_paragraph(doc, description, size=9.3)
                p.paragraph_format.left_indent = Inches(0.32)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if technologies:
                p = _add_paragraph(doc, f"Technologies: {technologies}", italic=True, size=9)
                p.paragraph_format.left_indent = Inches(0.32)
            if link:
                p = _add_paragraph(doc, link, size=8.5, color=ACCENT)
                p.paragraph_format.left_indent = Inches(0.32)

    if cv.get("certifications"):
        add_heading("Certifications")
        for item in cv.get("certifications"):
            if str(item).strip():
                _add_bullet(doc, item, size=9.5)

    if cv.get("languages"):
        add_heading("Languages")
        for item in cv.get("languages"):
            if str(item).strip():
                _add_bullet(doc, item, size=9.5)

    if cv.get("referees"):
        add_heading("Referees")
        for ref in cv.get("referees"):
            name = ref.get("name", "")
            position = ref.get("position") or ref.get("title", "")
            organization = ref.get("organization", "")
            email = ref.get("email", "")
            phone = ref.get("phone", "")

            parts = [p for p in [name, position, organization] if p]
            summary = " | ".join(parts)

            contacts = []
            if email:
                contacts.append(f"Email: {email}")
            if phone:
                contacts.append(f"Phone: {phone}")
            contact_line = " | ".join(contacts)

            text = summary
            if contact_line:
                text = f"{text} | {contact_line}" if text else contact_line

            if text:
                _add_bullet(doc, text, size=9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
