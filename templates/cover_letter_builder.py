import html
import re
from io import BytesIO

try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ModuleNotFoundError:
    docx = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    _DOCX_AVAILABLE = False


def cv_sender_address(cv: dict) -> str:
    lines = []
    if cv.get("location"):
        lines.append(str(cv.get("location", "")).strip())
    if cv.get("phone"):
        lines.append(f"Phone: {str(cv.get('phone', '')).strip()}")
    if cv.get("email"):
        lines.append(f"Email: {str(cv.get('email', '')).strip()}")
    return "\n".join([line for line in lines if line])


def default_cover_letter_data(cv: dict | None = None) -> dict:
    cv = cv or {}
    return {
        "name": str(cv.get("full_name", "")).strip(),
        "title": str(cv.get("headline", "")).strip(),
        "sender_address": cv_sender_address(cv),
        "recipient_address": "",
        "subject": "",
        "body": "",
        "signatory": str(cv.get("full_name", "")).strip(),
    }


def build_cover_letter_html(letter_data: dict) -> str:
    sender_address = str(letter_data.get("sender_address", ""))
    recipient_address = str(letter_data.get("recipient_address", ""))
    subject = html.escape(str(letter_data.get("subject", "")).strip())
    name = html.escape(str(letter_data.get("name", "")).strip())
    title = html.escape(str(letter_data.get("title", "")).strip())
    signatory = html.escape(str(letter_data.get("signatory", "")).strip())
    body = str(letter_data.get("body", ""))

    sender_address_html = "<br>".join(html.escape(line.strip()) for line in sender_address.splitlines() if line.strip())
    recipient_address_html = "<br>".join(html.escape(line.strip()) for line in recipient_address.splitlines() if line.strip())
    paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n", body.strip()) if segment.strip()]
    body_html = "".join(
        f"<p class='cover-letter-body'>{html.escape(paragraph).replace(chr(10), '<br>')}</p>" for paragraph in paragraphs
    )

    return f"""
    <div style="max-width: 860px; margin: 0 auto; background: #fff; border: 1px solid #d1d5db; padding: 28px; border-radius: 10px; font-family: Arial, sans-serif; color: #111827; line-height: 1.55;">
        <div style="margin-bottom: 20px;">{sender_address_html}</div>
        <div style="margin-bottom: 20px;">{recipient_address_html}</div>
        <div style="margin-bottom: 16px;"><strong>Subject:</strong> {subject}</div>
        {body_html}
        <div style="margin-top: 24px;">Sincerely,</div>
        <div style="margin-top: 56px; font-weight: 600;">{signatory or name}</div>
        <div style="margin-top: 2px; color: #4b5563;">{title}</div>
    </div>
    <style>
      .cover-letter-body {{
        margin: 0 0 14px 0;
        text-align: justify;
        text-justify: inter-word;
      }}
    </style>
    """


def build_cover_letter_text(letter_data: dict) -> str:
    sender_address = str(letter_data.get("sender_address", "")).strip()
    recipient_address = str(letter_data.get("recipient_address", "")).strip()
    subject = str(letter_data.get("subject", "")).strip()
    body = str(letter_data.get("body", "")).strip()
    signatory = str(letter_data.get("signatory", "")).strip() or str(letter_data.get("name", "")).strip()
    title = str(letter_data.get("title", "")).strip()

    lines = [sender_address, "", recipient_address, "", f"Subject: {subject}", "", body, "", "Sincerely,", "", signatory]
    if title:
        lines.append(title)
    return "\n".join(lines)


def build_cover_letter_docx(letter_data: dict) -> bytes:
    if not _DOCX_AVAILABLE:
        return b""

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    sender_address = str(letter_data.get("sender_address", "")).strip()
    recipient_address = str(letter_data.get("recipient_address", "")).strip()
    subject = str(letter_data.get("subject", "")).strip()
    body = str(letter_data.get("body", "")).strip()
    signatory = str(letter_data.get("signatory", "")).strip() or str(letter_data.get("name", "")).strip()
    title = str(letter_data.get("title", "")).strip()

    if sender_address:
        doc.add_paragraph(sender_address)
        doc.add_paragraph("")
    if recipient_address:
        doc.add_paragraph(recipient_address)
        doc.add_paragraph("")

    subject_para = doc.add_paragraph()
    subject_para.add_run("Subject: ").bold = True
    subject_para.add_run(subject)
    doc.add_paragraph("")

    body_paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n", body) if segment.strip()]
    for paragraph in body_paragraphs:
        p = doc.add_paragraph(paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("")
    doc.add_paragraph(signatory)
    if title:
        doc.add_paragraph(title)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
